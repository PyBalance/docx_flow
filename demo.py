#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Docx Toolkit - 代码即文档

本脚本是 `Docx Toolkit` 的终极功能演示。
它旨在成为 "代码即文档" 的典范：
1. 代码本身具有高度可读性，清晰地展示了所有功能的用法。
2. 脚本会自动生成一份Word文档，该文档既是所有功能的可视化成果展示，
   也是一份关于本工具包的说明书。

运行此脚本，即可生成 'demo_ultimate_output.docx' 文件并自动打开。
"""

import os
import subprocess
from docx import Document
from docx.shared import Cm, Inches
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

# ----------------------------------------------------------------------
# 导入 docx_flow 模块的各个组件
from docx_flow import DocxEditor
from docx_flow.conditions import (
    RegexCondition,
    TableTextCondition,
    FunctionCondition,
)
from docx_flow.actions import (
    AlignParagraphAction,
    ReplaceTextAction,
    SetTableWidthAction,
    RemoveTableBordersAction,
    AutoFitTableAction,
    SetFontSizeAction,
    SetTabStopAction,
    ClearAndSetTabStopAction,
    SetSectionOrientationAction,
    SetTableColumnWidthAction
)
# ----------------------------------------------------------------------


def create_comprehensive_demo_document(output_path="demo_ultimate_input.docx"):
    """
    创建一份结构化的、包含所有功能演示“原材料”的Word文档。
    """
    print(f"📝 正在创建综合演示文档: {output_path}...")
    
    doc = Document()
    s = doc.styles['Normal'].font
    s.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_heading('Docx Toolkit 终极功能演示', level=0)
    doc.add_paragraph(
        "本文件由Python脚本自动生成，用于演示Docx Toolkit的各项功能。脚本将加载本文档，"
        "并对下方标记为【目标】的元素执行操作，生成最终的输出文件。"
    )

    # --- 1. 段落操作 ---
    doc.add_heading('1. 段落 (Paragraph) 操作', level=1)

    doc.add_heading('1.1. 文本替换 (ReplaceTextAction)', level=2)
    doc.add_paragraph("【目标:文本替换】这段文本中的 '旧内容' 将被替换。")

    doc.add_heading('1.2. 对齐 (AlignParagraphAction)', level=2)
    doc.add_paragraph("【目标:居中对齐】这段文字将被居中。")
    doc.add_paragraph("【目标:右对齐】这段文字将被右对齐。")

    doc.add_heading('1.3. 字体大小 (SetFontSizeAction)', level=2)
    doc.add_paragraph("【目标:绝对字号】这段文字的字号将被设置为18磅。")
    doc.add_paragraph("【目标:相对字号】这段文字的字号将增加4磅。")

    doc.add_heading('1.4. 制表位 (TabStop Actions)', level=2)
    doc.add_paragraph("【目标:设置制表位】\t这行文字的制表位将被设置在8cm处。")
    doc.add_paragraph("【目标:清空并重设制表位】原有制表位将被清除，\t并重设在2cm处。")
    
    # --- 2. 表格操作 ---
    doc.add_heading('2. 表格 (Table) 操作', level=1)
    doc.add_paragraph("下面的一系列表格将分别展示不同的表格调整功能。")

    # 2.1 内容自适应
    doc.add_paragraph("\n--- 2.1 【目标:内容自适应】 (AutoFitTableAction) ---", style='Body Text')
    t1 = doc.add_table(rows=2, cols=3, style="Table Grid")
    t1.cell(0, 0).text = "标识: Content Autofit"
    t1.cell(0, 1).text = "短文本"
    t1.cell(0, 2).text = "这是一个非常非常长的文本，用于演示内容自适应效果"
    
    # 2.2 窗口自适应
    doc.add_paragraph("\n--- 2.2 【目标:窗口自适应】 (AutoFitTableAction) ---", style='Body Text')
    t2 = doc.add_table(rows=2, cols=3, style="Table Grid")
    t2.cell(0, 0).text = "标识: Window Autofit"

    # 2.3 固定宽度，列平分
    doc.add_paragraph("\n--- 2.3 【目标:固定宽度且列平分】 (AutoFitTableAction) ---", style='Body Text')
    t3 = doc.add_table(rows=2, cols=3, style="Table Grid")
    t3.cell(0, 0).text = "标识: Fixed (Equal)"

    # 2.4 首列按比例
    doc.add_paragraph("\n--- 2.4 【目标:首列按比例】 (AutoFitTableAction) ---", style='Body Text')
    t4 = doc.add_table(rows=2, cols=4, style="Table Grid")
    t4.cell(0, 0).text = "标识: Ratio"

    # 2.5 设置绝对宽度
    doc.add_paragraph("\n--- 2.5 【目标:绝对宽度】 (SetTableWidthAction) ---", style='Body Text')
    t5 = doc.add_table(rows=2, cols=2, style="Table Grid")
    t5.cell(0, 0).text = "标识: Absolute Width"

    # 2.6 设置各列精确宽度
    doc.add_paragraph("\n--- 2.6 【目标:各列精确宽度】 (SetTableColumnWidthAction) ---", style='Body Text')
    t6 = doc.add_table(rows=1, cols=3, style="Table Grid")
    t6.cell(0, 0).text = "标识: Column Widths"
    t6.cell(0, 1).text = "本列将为5cm"
    t6.cell(0, 2).text = "本列将为2cm"

    # 2.7 移除边框
    doc.add_paragraph("\n--- 2.7 【目标:移除边框】 (RemoveTableBordersAction) ---", style='Body Text')
    t7 = doc.add_table(rows=2, cols=2, style="Table Grid")
    t7.cell(0, 0).text = "标识: Border Removal"

    # --- 3. 节操作与高级筛选 ---
    doc.add_heading('3. 节操作与高级筛选', level=1)
    
    # 添加新的一节
    doc.add_section()
    doc.add_heading('3.1 页面方向 (SetSectionOrientationAction)', level=2)
    doc.add_paragraph("【目标:横向页面】本段落所在的整个“节”的页面方向将被设置为横向。")
    
    doc.add_heading('3.2 自定义函数筛选 (FunctionCondition)', level=2)
    p = doc.add_paragraph("【目标:自定义函数】")
    p.add_run("这个段落").bold = True
    p.add_run("由多个run组成，将被自定义函数匹配到。")

    doc.save(output_path)
    print("✅ 演示文档创建成功!")
    return output_path

def main():
    """主函数：创建文档 -> 链式调用处理 -> 保存结果"""
    print("🌟 Docx Toolkit 代码即文档终极演示 🌟")
    print("=" * 50)

    input_path = create_comprehensive_demo_document()
    output_path = "demo_ultimate_output.docx"

    print("\n🚀 开始使用 DocxEditor 对文档进行自动化处理...")
    editor = DocxEditor(input_path)

    # --- 1. 段落操作 ---
    print("\n--- 1. 处理段落 ---")
    # 1.1 文本替换
    editor.select_paragraphs()\
          .where(RegexCondition(r"【目标:文本替换】"))\
          .apply(ReplaceTextAction("旧内容", "焕然一新的内容"))\
          .apply(ReplaceTextAction("【目标:文本替换】", "【效果】"))
    print("  -> 1.1 文本替换... Done")

    # 1.2 对齐
    editor.select_paragraphs()\
          .where(RegexCondition(r"【目标:居中对齐】"))\
          .apply(AlignParagraphAction('center'))\
          .apply(ReplaceTextAction("【目标:居中对齐】", "【效果】"))
    editor.select_paragraphs()\
          .where(RegexCondition(r"【目标:右对齐】"))\
          .apply(AlignParagraphAction('right'))\
          .apply(ReplaceTextAction("【目标:右对齐】", "【效果】"))
    print("  -> 1.2 段落对齐... Done")

    # 1.3 字体大小
    editor.select_paragraphs()\
          .where(RegexCondition(r"【目标:绝对字号】"))\
          .apply(SetFontSizeAction(18))\
          .apply(ReplaceTextAction("【目标:绝对字号】", "【效果】"))
    editor.select_paragraphs()\
          .where(RegexCondition(r"【目标:相对字号】"))\
          .apply(SetFontSizeAction('+4'))\
          .apply(ReplaceTextAction("【目标:相对字号】", "【效果】"))
    print("  -> 1.3 字体大小调整... Done")
    
    # 1.4 制表位
    editor.select_paragraphs()\
          .where(RegexCondition(r"【目标:设置制表位】"))\
          .apply(SetTabStopAction(8.0))\
          .apply(ReplaceTextAction("【目标:设置制表位】", "【效果】"))
    editor.select_paragraphs()\
          .where(RegexCondition(r"【目标:清空并重设制表位】"))\
          .apply(ClearAndSetTabStopAction(2.0))\
          .apply(ReplaceTextAction("【目标:清空并重设制表位】", "【效果】"))
    print("  -> 1.4 制表位设置... Done")

    # --- 2. 表格操作 ---
    print("\n--- 2. 处理表格 ---")
    # 2.1 内容自适应
    editor.select_tables().where(TableTextCondition("Content Autofit"))\
          .apply(AutoFitTableAction('contents'))
    print("  -> 2.1 内容自适应... Done")
    # 2.2 窗口自适应
    editor.select_tables().where(TableTextCondition("Window Autofit"))\
          .apply(AutoFitTableAction('window'))
    print("  -> 2.2 窗口自适应... Done")
    # 2.3 固定宽度，列平分
    editor.select_tables().where(TableTextCondition("Fixed (Equal)"))\
          .apply(AutoFitTableAction('fixed'))
    print("  -> 2.3 固定宽度列平分... Done")
    # 2.4 首列按比例
    editor.select_tables().where(TableTextCondition("Ratio"))\
          .apply(AutoFitTableAction(first_col_ratio=0.5))
    print("  -> 2.4 首列50%比例... Done")
    # 2.5 设置绝对宽度
    editor.select_tables().where(TableTextCondition("Absolute Width"))\
          .apply(SetTableWidthAction(Cm(5)))
    print("  -> 2.5 绝对宽度5cm... Done")
    # 2.6 设置各列精确宽度
    editor.select_tables().where(TableTextCondition("Column Widths"))\
          .apply(SetTableColumnWidthAction([Inches(1.0), Cm(5), Cm(2)]))
    print("  -> 2.6 各列精确宽度... Done")
    # 2.7 移除边框
    editor.select_tables().where(TableTextCondition("Border Removal"))\
          .apply(RemoveTableBordersAction())
    print("  -> 2.7 移除边框... Done")

    # --- 3. 节操作与高级筛选 ---
    print("\n--- 3. 处理节与高级筛选 ---")
    # 3.1 页面方向
    # 筛选出在第1节（索引从0开始）的节对象并应用操作
    editor.select_sections().get_by_index(1).apply(SetSectionOrientationAction('landscape'))
    editor.select_paragraphs()\
        .where(RegexCondition(r"【目标:横向页面】"))\
        .apply(ReplaceTextAction("【目标:横向页面】", "【效果】"))
    print("  -> 3.1 页面方向设置为横向... Done")

    # 3.2 自定义函数筛选
    # 定义一个函数：检查段落是否包含超过2个run
    def has_multiple_runs(p: Paragraph) -> bool:
        return isinstance(p, Paragraph) and len(p.runs) > 2
    
    editor.select_paragraphs()\
          .where(FunctionCondition(has_multiple_runs))\
          .where(RegexCondition(r"【目标:自定义函数】"))\
          .apply(ReplaceTextAction("【目标:自定义函数】", "【效果】"))
    print("  -> 3.2 自定义函数筛选... Done")

    # --- 保存 ---
    editor.save(output_path)
    print("\n🎉 所有操作完成!")
    print(f"💾 最终文档已保存至: {output_path}")
    
    # --- 自动打开 ---
    try:
        os.startfile(os.path.realpath(output_path))
    except AttributeError:
        opener = "open" if os.name == "posix" else "xdg-open"
        subprocess.call([opener, os.path.realpath(output_path)])
    except Exception as e:
        print(f"(无法自动打开文件: {e})")

if __name__ == "__main__":
    main()