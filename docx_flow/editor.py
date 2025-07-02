# -*- coding: utf-8 -*-
"""
Docx Enhanced Toolkit - 编辑器模块

包含主编辑器类，负责文档的加载、保存和元素选择。
"""

from typing import Any

from docx import Document
from docx.document import Document as DocumentClass
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.section import Section

from .selector import FluentSelector


class DocxEditor:
    """库的主入口，负责文档的加载、保存和元素选择。"""
    def __init__(self, docx_path: str):
        self.document: DocumentClass = Document(docx_path)
        self._paragraph_section_map = {}
        self._table_section_map = {}
        self._build_element_section_map()

    def _build_element_section_map(self):
        """
        构建一个从元素索引到其节索引的映射。
        这是确定一个元素属于哪个节的更准确的方法。
        """
        self._paragraph_section_map = {}  # 映射段落索引 -> 节索引
        self._table_section_map = {}      # 映射表格索引 -> 节索引
        
        section_idx = 0
        para_idx = 0
        table_idx = 0

        # 遍历文档主体中的所有块级元素
        for block_element in self.document.element.body:
            # 检查元素的标签以确定是段落还是表格
            if block_element.tag.endswith('p'):  # 这是一个段落
                self._paragraph_section_map[para_idx] = section_idx
                
                # 检查此段落的属性中是否有分节符
                if block_element.pPr is not None and block_element.pPr.sectPr is not None:
                    section_idx += 1
                
                para_idx += 1

            elif block_element.tag.endswith('tbl'):  # 这是一个表格
                self._table_section_map[table_idx] = section_idx
                table_idx += 1

    def get_element_section_index(self, element: Any) -> int:
        """获取给定元素的节索引。"""
        if isinstance(element, Paragraph):
            # 通过其XML元素（这是稳定的）找到段落的索引。
            all_paras_xml = [p._p for p in self.document.paragraphs]
            try:
                para_index = all_paras_xml.index(element._p)
                return self._paragraph_section_map.get(para_index, 0)
            except ValueError:
                return 0  # 对于来自文档的元素，不应发生
        elif isinstance(element, Table):
            # 通过其XML元素找到表格的索引。
            all_tables_xml = [t._tbl for t in self.document.tables]
            try:
                table_index = all_tables_xml.index(element._tbl)
                return self._table_section_map.get(table_index, 0)
            except ValueError:
                return 0
        elif isinstance(element, Section):
            # Section对象是不同的，它们不是以相同方式成为主体的一部分。
            # 它们与 <w:sectPr> 元素相关联。
            # 节的列表应该是稳定的。
            try:
                # Sections list is not directly iterable in the same way, but we can create a list
                return list(self.document.sections).index(element)
            except ValueError:
                return 0
        return 0

    def select_paragraphs(self) -> FluentSelector:
        """选择文档中的所有段落。"""
        return FluentSelector(list(self.document.paragraphs), self)

    def select_tables(self) -> FluentSelector:
        """选择文档中的所有表格。"""
        return FluentSelector(list(self.document.tables), self)

    def select_sections(self) -> FluentSelector:
        """选择文档中的所有节。"""
        return FluentSelector(list(self.document.sections), self)

    def save(self, output_path: str):
        """保存修改后的文档。"""
        self.document.save(output_path)