# -*- coding: utf-8 -*-
"""
Docx Enhanced Toolkit - 核心模块

本模块包含了所有的核心类和功能实现。
"""

import re
from abc import ABC, abstractmethod
from typing import List, Any, Callable

# pip install python-docx
from docx import Document
from docx.document import Document as DocumentClass
from docx.section import Section, Sections
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# 条件类 (Conditions)
# =============================================================================

class Condition(ABC):
    """'条件'接口 (抽象基类)"""
    @abstractmethod
    def check(self, element: Any) -> bool:
        """
        检查给定的元素是否满足本条件。
        :param element: 文档元素 (Paragraph, Table, Section 等)。
        :return: 如果满足条件，返回 True，否则返回 False。
        """
        pass


class RegexCondition(Condition):
    """正则表达式条件：检查段落文本是否匹配特定模式。"""
    def __init__(self, pattern: str):
        self.pattern = re.compile(pattern)

    def check(self, element: Any) -> bool:
        if isinstance(element, Paragraph):
            return bool(self.pattern.search(element.text))
        return False


class TableColumnCondition(Condition):
    """表格列数条件：检查表格是否具有指定的列数。"""
    def __init__(self, column_count: int):
        self.column_count = column_count

    def check(self, element: Any) -> bool:
        if isinstance(element, Table):
            return len(element.columns) == self.column_count
        return False


class TableTextCondition(Condition):
    """表格文本条件：检查表格是否包含特定文本。"""
    def __init__(self, text: str):
        self.text = text

    def check(self, element: Any) -> bool:
        if isinstance(element, Table):
            for row in element.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self.text in paragraph.text:
                            return True
        return False


class FunctionCondition(Condition):
    """通用函数条件：使用一个自定义函数作为检查逻辑。"""
    def __init__(self, func: Callable[[Any], bool]):
        self.func = func

    def check(self, element: Any) -> bool:
        try:
            return self.func(element)
        except Exception:
            return False


# =============================================================================
# 操作类 (Actions)
# =============================================================================

class Action(ABC):
    """'操作'接口 (抽象基类)"""
    @abstractmethod
    def execute(self, element: Any) -> None:
        """
        对给定的元素执行具体操作。
        :param element: 文档元素 (Paragraph, Table, Section 等)。
        """
        pass


class RemoveTableBordersAction(Action):
    """移除表格边框的操作。将所有边设置为nil，防止Word回退到表格样式边框。"""
    def execute(self, element: Any) -> None:
        if not isinstance(element, Table):
            return
        # 获取表格的样式对象
        tbl = element._tbl
        tbl_pr = tbl.tblPr
        # 处理<w:tblBorders>
        if tbl_pr is not None:
            tbl_borders = tbl_pr.find(qn('w:tblBorders'))
            if tbl_borders is None:
                tbl_borders = OxmlElement('w:tblBorders')
                tbl_pr.append(tbl_borders)
            for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                tag = qn(f"w:{edge}")
                border = tbl_borders.find(tag)
                if border is None:
                    border = OxmlElement(f"w:{edge}")
                    tbl_borders.append(border)
                border.set(qn('w:val'), 'nil')
        # 为所有单元格移除边框（同样设置为nil）
        for row in element.rows:
            for cell in row.cells:
                tc = cell._tc
                tc_pr = tc.tcPr
                if tc_pr is not None:
                    tc_borders = tc_pr.find(qn('w:tcBorders'))
                    if tc_borders is None:
                        tc_borders = OxmlElement('w:tcBorders')
                        tc_pr.append(tc_borders)
                    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                        tag = qn(f"w:{edge}")
                        border = tc_borders.find(tag)
                        if border is None:
                            border = OxmlElement(f"w:{edge}")
                            tc_borders.append(border)
                        border.set(qn('w:val'), 'nil')


class SetTableWidthAction:
    def __init__(self, width):
        # 接受任何 Length 类型，也可直接给 int(twips)
        self.width = width

    def execute(self, element: Any) -> None:
        if not isinstance(element, Table):
            return

        # 关闭自动调整，保持固定列宽
        element.allow_autofit = False

        tbl_pr = element._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            element._tbl.insert(0, tbl_pr)

        # 固定布局
        layout = tbl_pr.find(qn('w:tblLayout'))
        if layout is None:
            layout = OxmlElement('w:tblLayout')
            layout.set(qn('w:type'), 'fixed')
            tbl_pr.append(layout)

        # 表格宽度
        tbl_w = tbl_pr.find(qn('w:tblW'))
        if tbl_w is None:
            tbl_w = OxmlElement('w:tblW')
            tbl_pr.append(tbl_w)

        # 把输入宽度统一转为 twips
        length = self.width
        if hasattr(length, 'twips'):
            length = length.twips      # Inches/Cm/Pt 等
        tbl_w.set(qn('w:w'), str(int(length)))
        tbl_w.set(qn('w:type'), 'dxa')

class AutoFitTableAction(Action):
    """自动调整表格大小的操作。"""
    def __init__(self, autofit_type: str = 'contents', first_col_ratio: float = None):
        """
        :param autofit_type: 自动调整类型 (当 first_col_ratio 未指定时生效)
            - 'contents': 根据内容自动调整
            - 'window': 自动调整到窗口宽度
            - 'fixed': 固定列宽 (平均分配)
        :param first_col_ratio: 一个0到1之间的浮点数，用于指定第一列的宽度占总宽度的比例。
                                 如果设置了此参数，表格总宽将适应窗口，并按比例分配列宽，
                                 此时 autofit_type 会被忽略。
                                 例如, 0.67 表示首列占67%宽度，其余列平分33%的宽度。
        """
        if first_col_ratio is not None and not (0 < first_col_ratio < 1):
            raise ValueError("first_col_ratio 必须是一个大于0且小于1的浮点数。")
            
        self.autofit_type = autofit_type.lower()
        self.first_col_ratio = first_col_ratio

    def execute(self, element: Any) -> None:
        if not isinstance(element, Table):
            return
        
        
        tbl = element._tbl
        # 恢复为您原来的、正确的方法来获取或创建 tblPr
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)

        # 清理可能冲突的旧设置
        for tag in ['w:tblLayout', 'w:tblW']:
            existing = tbl_pr.find(qn(tag))
            if existing is not None:
                tbl_pr.remove(existing)

        # 优先处理 first_col_ratio 逻辑
        if self.first_col_ratio is not None:
            col_count = len(element.columns)
            if col_count == 0:
                return

            # 1. 设置为固定布局，这是按比例分配宽度的前提
            tbl_layout = OxmlElement('w:tblLayout')
            tbl_layout.set(qn('w:type'), 'fixed')
            tbl_pr.append(tbl_layout)

            # 2. 设置表格总宽度为100%窗口宽度 (参考 _apply_autofit_window)
            tbl_w = OxmlElement('w:tblW')
            tbl_w.set(qn('w:w'), '5000')  # 5000 在OOXML中代表 100%
            tbl_w.set(qn('w:type'), 'pct')
            tbl_pr.append(tbl_w)
            
            # 关闭高阶API的自动调整
            element.allow_autofit = False

            # 3. 计算并以百分比形式设置每一列的宽度
            # Word中，列宽通常由第一行单元格的宽度定义
            if col_count > 1:
                # 在pct单位下，总宽度是5000
                first_col_pct = int(self.first_col_ratio * 5000)
                remaining_pct = 5000 - first_col_pct
                other_col_pct = int(remaining_pct / (col_count - 1))
                widths = [first_col_pct] + [other_col_pct] * (col_count - 1)
            else: # 如果只有一列
                widths = [5000]

            # 直接操作第一行单元格的XML来设置百分比宽度
            first_row_cells = element.rows[0].cells
            for i, cell in enumerate(first_row_cells):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.get_or_add_tcW()
                tc_w.set(qn('w:w'), str(widths[i]))
                tc_w.set(qn('w:type'), 'pct')

        # 如果没有设置 first_col_ratio，则执行原有的逻辑
        else:
            if self.autofit_type == 'contents':
                element.allow_autofit = True
                tbl_layout = OxmlElement('w:tblLayout')
                tbl_layout.set(qn('w:type'), 'autofit')
                tbl_pr.append(tbl_layout)
                tbl_w = OxmlElement('w:tblW')
                tbl_w.set(qn('w:w'), '0')
                tbl_w.set(qn('w:type'), 'auto')
                tbl_pr.append(tbl_w)
                
            elif self.autofit_type == 'window':
                element.allow_autofit = True
                tbl_layout = OxmlElement('w:tblLayout')
                tbl_layout.set(qn('w:type'), 'autofit')
                tbl_pr.append(tbl_layout)
                tbl_w = OxmlElement('w:tblW')
                tbl_w.set(qn('w:w'), '5000')
                tbl_w.set(qn('w:type'), 'pct')
                tbl_pr.append(tbl_w)
                
            elif self.autofit_type == 'fixed':
                element.allow_autofit = False
                tbl_layout = OxmlElement('w:tblLayout')
                tbl_layout.set(qn('w:type'), 'fixed')
                tbl_pr.append(tbl_layout)
                
                total_width = Inches(6.5)
                col_count = len(element.columns)
                if col_count > 0:
                    col_width = int(total_width / col_count)
                    tbl_w = OxmlElement('w:tblW')
                    tbl_w.set(qn('w:w'), str(int(total_width.twips)))
                    tbl_w.set(qn('w:type'), 'dxa')
                    tbl_pr.append(tbl_w)
                    for column in element.columns:
                        column.width = col_width


    def _apply_autofit_contents(self, tbl_pr):
        """设置为根据内容自动调整"""
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_layout.set(qn('w:type'), 'autofit')
        tbl_pr.append(tbl_layout)
        
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'), '0')
        tbl_w.set(qn('w:type'), 'auto')
        tbl_pr.append(tbl_w)

    def _apply_autofit_window(self, tbl_pr):
        """设置为根据窗口宽度自动调整"""
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_layout.set(qn('w:type'), 'autofit')
        tbl_pr.append(tbl_layout)
        
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'), '5000')  # 5000 = 100% (in 50ths of a percent)
        tbl_w.set(qn('w:type'), 'pct')
        tbl_pr.append(tbl_w)

    def _apply_fixed_layout(self, tbl_pr, table):
        """设置为固定宽度并平均分配列宽"""
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_layout.set(qn('w:type'), 'fixed')
        tbl_pr.append(tbl_layout)
        
        total_width = Inches(6.5)
        col_count = len(table.columns)
        if col_count > 0:
            col_width = int(total_width / col_count)
            
            tbl_w = OxmlElement('w:tblW')
            tbl_w.set(qn('w:w'), str(total_width.twips))
            tbl_w.set(qn('w:type'), 'dxa')
            tbl_pr.append(tbl_w)
            
            for column in table.columns:
                column.width = col_width
            

class SetTableColumnWidthAction:
    """
    将表格切换到“固定布局”并精确设置每列宽度。
    """

    def __init__(self, widths):
        """
        :param widths: 与表格列数等长的宽度序列，元素可用 Inches()/Cm()/Pt() 等。
        """
        self.widths = widths

    # ──────────────────────────────────────────
    def execute(self, table: Table) -> None:
        if not isinstance(table, Table):
            raise TypeError("只能传入 python-docx.Table 对象")

        if len(table.columns) != len(self.widths):
            raise ValueError(
                f"表格有 {len(table.columns)} 列，但传入了 {len(self.widths)} 个宽度值"
            )

        # ① 关闭自动调整，固定布局
        if hasattr(table, "autofit"):           # ≥0.8.11
            table.autofit = False
        else:
            table.allow_autofit = False         # 旧版别名

        # ② 为每个单元格写入宽度
        for row in table.rows:
            for idx, width in enumerate(self.widths):
                row.cells[idx].width = width

class AlignParagraphAction(Action):
    """段落对齐的操作。"""
    def __init__(self, alignment: str):
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        self.alignment = alignment_map.get(alignment.lower())

    def execute(self, element: Any) -> None:
        if isinstance(element, Paragraph) and self.alignment is not None:
            element.alignment = self.alignment


class SetTabStopAction(Action):
    """设置段落制表位的操作。"""
    def __init__(self, position_in_cm: float):
        self.position = Cm(position_in_cm)

    def execute(self, element: Any) -> None:
        if isinstance(element, Paragraph):
            p_format = element.paragraph_format
            p_format.tab_stops.add_tab_stop(self.position)

class ClearAndSetTabStopAction(Action):
    """清除现有制表位并设置新制表位的操作。"""
    def __init__(self, position_in_cm: float):
        self.position = Cm(position_in_cm)

    def execute(self, element: Any) -> None:
        if isinstance(element, Paragraph):
            p_format = element.paragraph_format
            # 清除所有现有制表位
            p_format.tab_stops.clear_all()
            # 设置新的制表位
            p_format.tab_stops.add_tab_stop(self.position)

class ReplaceTextAction(Action):
    """替换文本的操作（段落和表格内）。"""
    def __init__(self, old_text: str, new_text: str):
        self.old_text = old_text
        self.new_text = new_text

    def execute(self, element: Any) -> None:
        if isinstance(element, Paragraph):
            self.replace_in_paragraph(element)
        elif isinstance(element, Table):
            for row in element.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self.replace_in_paragraph(p)
    
    def replace_in_paragraph(self, paragraph: Paragraph):
        """在段落中替换文本，保留格式。"""
        for run in paragraph.runs:
            if self.old_text in run.text:
                run.text = run.text.replace(self.old_text, self.new_text)


class SetFontSizeAction(Action):
    """修改字号的操作。"""
    def __init__(self, size: Any):
        """
        :param size: 字号。可以是：
            - 整数或浮点数 (e.g., 12) 来设置绝对字号 (单位: Pt)。
            - 字符串 (e.g., '+2', '-1') 来相对调整字号。
        """
        self.absolute_size = None
        self.relative_change = None

        if isinstance(size, str):
            try:
                if size.startswith('+'):
                    self.relative_change = int(size[1:])
                elif size.startswith('-'):
                    self.relative_change = -int(size[1:])
                else:
                    self.absolute_size = Pt(int(size))
            except (ValueError, TypeError):
                raise ValueError("Invalid size format for string. Use '+n', '-n', or a number string.")
        elif isinstance(size, (int, float)):
            self.absolute_size = Pt(size)
        else:
            raise TypeError("Size must be an int, float, or string.")

    def execute(self, element: Any) -> None:
        if isinstance(element, Paragraph):
            self._apply_to_paragraph(element)
        elif isinstance(element, Table):
            for row in element.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._apply_to_paragraph(p)

    def _apply_to_paragraph(self, paragraph: Paragraph):
        """将字号更改应用于段落中的所有 run。"""
        for run in paragraph.runs:
            font = run.font
            if self.absolute_size:
                font.size = self.absolute_size
            elif self.relative_change:
                # 只有在已明确设置了字号的情况下才能进行相对更改
                current_size = font.size
                if current_size is not None:
                    new_size_pt = current_size.pt + self.relative_change
                    if new_size_pt > 0:  # 字号必须是正数
                        font.size = Pt(new_size_pt)


class SetSectionOrientationAction(Action):
    """设置节页面方向的操作。"""
    def __init__(self, orientation: str = 'landscape'):
        self.orientation = WD_ORIENT.LANDSCAPE if orientation.lower() == 'landscape' else WD_ORIENT.PORTRAIT
    
    def execute(self, element: Any) -> None:
        if not isinstance(element, Section):
            raise TypeError("只能对 Section 对象使用 SetSectionOrientationAction")
        
        original_width = element.page_width
        original_height = element.page_height
            
        print(f"设置节页面方向为: {self.orientation.name}", f"原始宽度: {original_width}, 高度: {original_height}")
        # 检查是否有有效的页面尺寸
        if original_width is None or original_height is None:
            print("警告: 节没有设置页面尺寸，使用默认 A4 尺寸。")
            # 使用A4纸张尺寸作为默认值 (单位: Twips, 1 inch = 1440 twips)
            # A4: 210mm x 297mm = 8.27" x 11.69"
            from docx.shared import Inches
            original_width = Inches(8.27)
            original_height = Inches(11.69)
            
        # 设为横向
        element.orientation = WD_ORIENT.LANDSCAPE
        # 交换宽高，确保宽度大于高度（横向特征）
        element.page_width = max(original_width, original_height)
        element.page_height = min(original_width, original_height)




# =============================================================================
# 流畅选择器 (Fluent Selector)
# =============================================================================

class FluentSelector:
    """流畅选择器，支持链式调用。"""
    def __init__(self, elements: List[Any], editor: 'DocxEditor'):
        self._elements = elements
        self._editor = editor

    def get_by_index(self, index: int) -> 'FluentSelector':
        # 把负索引转换为正索引
        if index < 0:
            index += len(self._elements)
        if 0 <= index < len(self._elements):
            return FluentSelector([self._elements[index]], self._editor)
        return FluentSelector([], self._editor)

    def where(self, condition: Condition) -> 'FluentSelector':
        """根据 Condition 对象筛选元素。"""
        filtered = [elem for elem in self._elements if condition.check(elem)]
        return FluentSelector(filtered, self._editor)

    def in_section(self, section_index: int) -> 'FluentSelector':
        """按节索引筛选元素的便捷方法。"""
        def check_func(element):
            return self._editor.get_element_section_index(element) == section_index
        return self.where(FunctionCondition(check_func))
    
    def from_section(self, start_section_index: int) -> 'FluentSelector':
        """从指定节开始筛选元素的便捷方法。"""
        def check_func(element):
            idx = self._editor.get_element_section_index(element)
            return idx is not None and idx >= start_section_index
        return self.where(FunctionCondition(check_func))

    def apply(self, action: Action) -> 'FluentSelector':
        """将一个 Action 应用于所有当前选中的元素。"""
        if not self._elements:
            print("没有选中任何元素，无法执行操作。")
            
        for element in self._elements:
            action.execute(element)
        return self

    @property
    def count(self) -> int:
        """返回当前选中元素的数量。"""
        return len(self._elements)

    def get(self) -> List[Any]:
        """获取所有当前选中的元素。"""
        return self._elements


# =============================================================================
# 主入口类 (DocxEditor)
# =============================================================================

class DocxEditor:
    """库的主入口，负责文档的加载、保存和元素选择。"""
    def __init__(self, docx_path: str):
        self.document: DocumentClass = Document(docx_path)
        self._paragraph_section_map = {}
        self._table_section_map = {}
        self._build_element_section_map()

    def _build_element_section_map(self):
        """
        构建一个从元素索引到其节索引的映射。
        这是确定一个元素属于哪个节的更准确的方法。
        """
        self._paragraph_section_map = {}  # 映射段落索引 -> 节索引
        self._table_section_map = {}      # 映射表格索引 -> 节索引
        
        section_idx = 0
        para_idx = 0
        table_idx = 0

        # 遍历文档主体中的所有块级元素
        for block_element in self.document.element.body:
            # 检查元素的标签以确定是段落还是表格
            if block_element.tag.endswith('p'):  # 这是一个段落
                self._paragraph_section_map[para_idx] = section_idx
                
                # 检查此段落的属性中是否有分节符
                if block_element.pPr is not None and block_element.pPr.sectPr is not None:
                    section_idx += 1
                
                para_idx += 1

            elif block_element.tag.endswith('tbl'):  # 这是一个表格
                self._table_section_map[table_idx] = section_idx
                table_idx += 1

    def get_element_section_index(self, element: Any) -> int:
        """获取给定元素的节索引。"""
        if isinstance(element, Paragraph):
            # 通过其XML元素（这是稳定的）找到段落的索引。
            all_paras_xml = [p._p for p in self.document.paragraphs]
            try:
                para_index = all_paras_xml.index(element._p)
                return self._paragraph_section_map.get(para_index, 0)
            except ValueError:
                return 0  # 对于来自文档的元素，不应发生
        elif isinstance(element, Table):
            # 通过其XML元素找到表格的索引。
            all_tables_xml = [t._tbl for t in self.document.tables]
            try:
                table_index = all_tables_xml.index(element._tbl)
                return self._table_section_map.get(table_index, 0)
            except ValueError:
                return 0
        elif isinstance(element, Section):
            # Section对象是不同的，它们不是以相同方式成为主体的一部分。
            # 它们与 <w:sectPr> 元素相关联。
            # 节的列表应该是稳定的。
            try:
                # Sections list is not directly iterable in the same way, but we can create a list
                return list(self.document.sections).index(element)
            except ValueError:
                return 0
        return 0

    def select_paragraphs(self) -> FluentSelector:
        """选择文档中的所有段落。"""
        return FluentSelector(list(self.document.paragraphs), self)

    def select_tables(self) -> FluentSelector:
        """选择文档中的所有表格。"""
        return FluentSelector(list(self.document.tables), self)

    def select_sections(self) -> FluentSelector:
        """选择文档中的所有节。"""
        return FluentSelector(list(self.document.sections), self)

    def save(self, output_path: str):
        """保存修改后的文档。"""
        self.document.save(output_path)
