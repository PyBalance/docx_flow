#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
高级功能示例 - Docx Flow

演示更复杂的文档处理场景，包括：
- 多节文档处理
- 自定义条件函数
- 表格高级操作
- 批量格式化
"""

from docx import Document
from docx.shared import Inches, Cm
from docx_flow import DocxEditor
from docx_flow.conditions import (
    RegexCondition, 
    TableColumnCondition, 
    FunctionCondition
)
from docx_flow.actions import (
    ReplaceTextAction,
    AlignParagraphAction,
    SetFontSizeAction,
    SetTableColumnWidthAction,
    RemoveTableBordersAction,
    SetSectionOrientationAction
)


def create_complex_document():
    """创建一个复杂的多节文档"""
    doc = Document()
    
    # 第一节：封面
    doc.add_heading('年度工作报告', 0)
    doc.add_paragraph('机密文档 - 仅供内部使用')
    doc.add_paragraph('编制日期：2024年')
    
    # 第二节：主要内容 
    doc.add_section()
    doc.add_heading('第一章 业务概况', 1)
    doc.add_paragraph('重点关注：本章节包含重要的业务数据分析。')
    doc.add_paragraph('这是一个很长的段落，包含了大量的业务分析内容，需要进行特殊的格式化处理。字数超过了50个字符，应该被识别为长段落。')
    doc.add_paragraph('短段落')
    
    # 添加数据表格
    table1 = doc.add_table(rows=4, cols=4)
    table1.style = 'Table Grid'
    table1.cell(0, 0).text = '业务数据'
    table1.cell(0, 1).text = '2022'
    table1.cell(0, 2).text = '2023' 
    table1.cell(0, 3).text = '2024'
    
    # 第三节：附录
    doc.add_section()
    doc.add_heading('附录', 1)
    doc.add_paragraph('附录内容：详细的技术说明文档。')
    
    # 添加另一个表格
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = 'Table Grid'
    table2.cell(0, 0).text = '技术参数'
    table2.cell(0, 1).text = '数值'
    
    doc.save('complex_input.docx')
    print("✅ 复杂示例文档已创建: complex_input.docx")


def main():
    """主函数：演示高级功能"""
    print("🚀 Docx Flow 高级功能示例")
    print("=" * 40)
    
    # 1. 创建复杂文档
    create_complex_document()
    
    # 2. 使用 DocxEditor 处理
    editor = DocxEditor('complex_input.docx')
    
    print("\n📝 开始高级处理...")
    
    # 3. 按节处理：只处理第一节的内容
    print("  → 处理第一节（封面）...")
    editor.select_paragraphs() \
        .in_section(0) \
        .where(RegexCondition(r'机密')) \
        .apply(ReplaceTextAction('机密文档', '🔒 机密文档')) \
        .apply(AlignParagraphAction('center'))
    
    # 4. 使用自定义函数条件：处理长段落
    print("  → 识别并处理长段落...")
    def is_long_paragraph(p):
        """自定义条件：段落长度超过50字符"""
        from docx.text.paragraph import Paragraph
        return isinstance(p, Paragraph) and len(p.text.strip()) > 50
    
    editor.select_paragraphs() \
        .where(FunctionCondition(is_long_paragraph)) \
        .apply(AlignParagraphAction('justify')) \
        .apply(SetFontSizeAction(11))
    
    # 5. 处理重点内容
    print("  → 标记重点内容...")
    editor.select_paragraphs() \
        .where(RegexCondition(r'重点关注')) \
        .apply(ReplaceTextAction('重点关注', '🎯 重点关注')) \
        .apply(SetFontSizeAction(13))
    
    # 6. 表格高级操作：4列表格设置精确列宽
    print("  → 设置4列表格的精确列宽...")
    column_widths = [Cm(3), Cm(2.5), Cm(2.5), Cm(2.5)]
    editor.select_tables() \
        .where(TableColumnCondition(4)) \
        .apply(SetTableColumnWidthAction(column_widths))
    
    # 7. 移除2列表格的边框
    print("  → 移除2列表格的边框...")
    editor.select_tables() \
        .where(TableColumnCondition(2)) \
        .apply(RemoveTableBordersAction())
    
    # 8. 从第二节开始的内容添加水印文字
    print("  → 处理第二节及以后的内容...")
    editor.select_paragraphs() \
        .from_section(1) \
        .where(RegexCondition(r'附录')) \
        .apply(ReplaceTextAction('附录', '📎 附录'))
    
    # 9. 设置第三节为横向页面
    print("  → 设置第三节为横向页面...")
    editor.select_sections() \
        .get_by_index(2) \
        .apply(SetSectionOrientationAction('landscape'))
    
    # 10. 保存结果
    editor.save('complex_output.docx')
    
    print("\n✅ 高级处理完成！")
    print("📄 输出文件: complex_output.docx")
    print("\n🔍 处理效果：")
    print("  • 第一节机密文档标记为🔒，居中显示")
    print("  • 长段落（>50字符）自动两端对齐，字号11")  
    print("  • 重点关注标记为🎯，字号13")
    print("  • 4列表格设置精确列宽（3cm, 2.5cm, 2.5cm, 2.5cm）")
    print("  • 2列表格移除边框")
    print("  • 附录标记为📎")
    print("  • 第三节设置为横向页面")


if __name__ == "__main__":
    main()