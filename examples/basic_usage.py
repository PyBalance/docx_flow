#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
基本使用示例 - Docx Flow

演示如何使用 Docx Flow 进行基本的文档操作。
"""

from docx import Document
from docx_flow import DocxEditor
from docx_flow.conditions import RegexCondition, TableTextCondition
from docx_flow.actions import (
    ReplaceTextAction, 
    AlignParagraphAction, 
    SetFontSizeAction,
    AutoFitTableAction
)


def create_sample_document():
    """创建一个示例文档用于演示"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('公司季度报告', 0)
    
    # 添加段落
    doc.add_paragraph('重要提示：本报告包含机密信息。')
    doc.add_paragraph('注意：请仔细阅读以下内容。')
    doc.add_paragraph('普通段落：这是一段普通的文字内容。')
    
    # 添加表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # 表格标识
    table.cell(0, 0).text = '财务数据'
    table.cell(0, 1).text = 'Q1'
    table.cell(0, 2).text = 'Q2'
    table.cell(1, 0).text = '收入'
    table.cell(1, 1).text = '100万'
    table.cell(1, 2).text = '120万'
    table.cell(2, 0).text = '支出'
    table.cell(2, 1).text = '80万'
    table.cell(2, 2).text = '90万'
    
    doc.save('sample_input.docx')
    print("✅ 示例文档已创建: sample_input.docx")


def main():
    """主函数：演示基本用法"""
    print("🚀 Docx Flow 基本使用示例")
    print("=" * 40)
    
    # 1. 创建示例文档
    create_sample_document()
    
    # 2. 使用 DocxEditor 处理文档
    editor = DocxEditor('sample_input.docx')
    
    print("\n📝 开始处理文档...")
    
    # 3. 文本替换和格式化
    print("  → 处理重要提示...")
    editor.select_paragraphs() \
        .where(RegexCondition(r'重要提示')) \
        .apply(ReplaceTextAction('重要提示', '⭐ 重要提示')) \
        .apply(AlignParagraphAction('center')) \
        .apply(SetFontSizeAction(14))
    
    # 4. 处理注意事项
    print("  → 处理注意事项...")
    editor.select_paragraphs() \
        .where(RegexCondition(r'注意：')) \
        .apply(ReplaceTextAction('注意：', '⚠️ 注意：')) \
        .apply(SetFontSizeAction(12))
    
    # 5. 处理表格
    print("  → 优化表格显示...")
    editor.select_tables() \
        .where(TableTextCondition('财务数据')) \
        .apply(AutoFitTableAction('window'))
    
    # 6. 保存结果
    editor.save('sample_output.docx')
    
    print("\n✅ 处理完成！")
    print("📄 输出文件: sample_output.docx")
    print("\n🔍 处理效果：")
    print("  • 重要提示添加了星号标记，居中显示，字号14")
    print("  • 注意事项添加了警告图标，字号12")
    print("  • 财务数据表格调整为窗口自适应")


if __name__ == "__main__":
    main()