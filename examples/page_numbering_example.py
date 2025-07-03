#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
页码功能示例

演示如何使用 docx_flow 为 Word 文档添加和管理页码。
"""

from docx import Document
from docx_flow import DocxEditor
from docx_flow.actions import AddPageNumberAction, ClearPageNumberAction


def create_test_document():
    """创建测试文档"""
    doc = Document()
    
    # 第一节：封面页
    doc.add_heading('封面页', level=0)
    doc.add_paragraph('这是文档的封面页，不应该有页码。')
    
    # 第二节：目录
    doc.add_section()
    doc.add_heading('目录', level=0)
    doc.add_paragraph('第一章 ............................ 1')
    doc.add_paragraph('第二章 ............................ 5')
    doc.add_paragraph('附录 .............................. 10')
    
    # 第三节：第一章
    doc.add_section()
    doc.add_heading('第一章 引言', level=1)
    doc.add_paragraph('这是第一章的内容。从这里开始应该有页码。')
    for i in range(3):
        doc.add_paragraph(f'第一章第{i+1}段内容。' * 15)
    
    # 第四节：第二章
    doc.add_section()
    doc.add_heading('第二章 方法', level=1)
    doc.add_paragraph('这是第二章的内容，页码应该继续。')
    for i in range(3):
        doc.add_paragraph(f'第二章第{i+1}段内容。' * 15)
    
    # 第五节：附录
    doc.add_section()
    doc.add_heading('附录', level=1)
    doc.add_paragraph('这是附录内容，通常重新开始编页码。')
    for i in range(2):
        doc.add_paragraph(f'附录第{i+1}段内容。' * 15)
    
    filename = "test_document.docx"
    doc.save(filename)
    print(f"创建测试文档: {filename} (包含 {len(doc.sections)} 个节)")
    return filename


def main():
    """主函数"""
    print("页码功能测试")
    print("=" * 30)
    
    # 创建测试文档
    input_file = create_test_document()
    
    try:
        editor = DocxEditor(input_file)
        print(f"文档包含 {editor.select_sections().count} 个节")
        
        # 测试1: 清除所有页码
        print("\n测试1: 清除所有页码")
        editor.select_sections().apply(ClearPageNumberAction())
        
        # 测试2: 为第3节添加页码（从1开始）
        print("测试2: 为第3节添加页码（从1开始）")
        editor.select_sections().get_by_index(2).apply(AddPageNumberAction(
            start_number=1, 
            restart_numbering=True
        ))
        
        # 测试3: 为第4节添加连续页码
        print("测试3: 为第4节添加连续页码")
        editor.select_sections().get_by_index(3).apply(AddPageNumberAction(
            restart_numbering=False
        ))
        
        # 测试4: 为第5节重新开始页码
        print("测试4: 为第5节重新开始页码")
        editor.select_sections().get_by_index(4).apply(AddPageNumberAction(
            start_number=1,
            restart_numbering=True
        ))
        
        # 保存结果
        output_file = "test_output.docx"
        editor.save(output_file)
        print(f"\n测试完成，已保存: {output_file}")
        print("预期结果:")
        print("  - 第1节（封面）：无页码")
        print("  - 第2节（目录）：无页码")
        print("  - 第3节（第一章）：页码 1")
        print("  - 第4节（第二章）：页码 2")
        print("  - 第5节（附录）：页码 1")
        print("\n请在Word中打开文档，按F9更新字段验证")
        
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()